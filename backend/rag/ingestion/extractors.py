"""Step 1 of the RAG pipeline: get plain text out of an uploaded file."""

from __future__ import annotations

import csv
import io
from pathlib import Path

from app.core.errors import ValidationError
from app.core.logging_config import get_logger

log = get_logger("rag.extract")

MAX_PDF_PAGES = 400
MAX_EXTRACTED_CHARS = 800_000


class ExtractionResult:
    def __init__(self, text: str, *, pages: int = 0, note: str = ""):
        self.text = text
        self.pages = pages
        self.note = note


def _extract_pdf(data: bytes) -> ExtractionResult:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValidationError(
            "PDF support requires `pypdf`. Install it or upload a .txt/.md file."
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - malformed uploads are user error
        raise ValidationError("That PDF could not be read. It may be corrupt or encrypted.") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001
            raise ValidationError("That PDF is password protected.") from exc

    pages = reader.pages[:MAX_PDF_PAGES]
    parts: list[str] = []
    for index, page in enumerate(pages):
        try:
            content = page.extract_text() or ""
        except Exception:  # noqa: BLE001 - skip unreadable pages, keep the rest
            log.warning("Skipped unreadable PDF page %s", index + 1)
            continue
        if content.strip():
            parts.append(f"[page {index + 1}]\n{content}")

    note = ""
    if len(reader.pages) > MAX_PDF_PAGES:
        note = f"Only the first {MAX_PDF_PAGES} of {len(reader.pages)} pages were indexed."

    return ExtractionResult("\n\n".join(parts), pages=len(pages), note=note)


def _extract_docx(data: bytes) -> ExtractionResult:
    try:
        import docx
    except ImportError as exc:
        raise ValidationError(
            "DOCX support requires `python-docx`. Install it or upload a .txt/.md file."
        ) from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ValidationError("That Word document could not be read.") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return ExtractionResult("\n".join(parts))


def _extract_csv(data: bytes) -> ExtractionResult:
    text = data.decode("utf-8", errors="replace")
    try:
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)[:2000]
    except csv.Error:
        return ExtractionResult(text)

    if not rows:
        return ExtractionResult("")

    header, *body = rows
    lines = []
    for row in body:
        pairs = [
            f"{(header[i] if i < len(header) else f'col{i}').strip()}: {value.strip()}"
            for i, value in enumerate(row)
            if value.strip()
        ]
        if pairs:
            lines.append("; ".join(pairs))
    return ExtractionResult("\n".join(lines))


def _extract_plain(data: bytes) -> ExtractionResult:
    return ExtractionResult(data.decode("utf-8", errors="replace"))


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".csv": _extract_csv,
    ".txt": _extract_plain,
    ".md": _extract_plain,
    ".markdown": _extract_plain,
}

SUPPORTED_EXTENSIONS = frozenset(_EXTRACTORS)


def extract_text(filename: str, data: bytes) -> ExtractionResult:
    suffix = Path(filename).suffix.lower()
    extractor = _EXTRACTORS.get(suffix)
    if not extractor:
        raise ValidationError(
            f"Unsupported file type '{suffix or 'unknown'}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    result = extractor(data)
    if len(result.text) > MAX_EXTRACTED_CHARS:
        result.text = result.text[:MAX_EXTRACTED_CHARS]
        result.note = (result.note + " Document truncated at 800k characters.").strip()

    if not result.text.strip():
        raise ValidationError(
            "No readable text was found in that file. Scanned PDFs need OCR before upload."
        )
    return result
